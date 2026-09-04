"""Client-facing incident report for the 2026-09-04 late-payment-webhook incident.

Usage: python scripts/create-incident-report-2026-09-04.py
Writes docs/Wolastoq Bingo - Payment Incident Report 2026-09-04.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Wolastoq Bingo - Payment Incident Report 2026-09-04.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GREY = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xB9, 0x1C, 0x1C)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def table(doc, headers, rows, widths=None, header_fill="1A3A5C"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, header_fill)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)
            if r_i % 2 == 1:
                shade(cells[i], "F3F4F6")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title block ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("Wolastoq Bingo Online Booking")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    r.bold = True
    t2 = doc.add_paragraph()
    r = t2.add_run("Incident Report: Payments Taken Without a Confirmed Reservation")
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = NAVY
    meta = doc.add_paragraph()
    meta.add_run(
        "Incident date: Thursday, September 4, 2026    |    Report date: September 5, 2026\n"
        "Prepared for: Saint Mary's Entertainment Centre / Wolastoq Casino management\n"
        "Prepared by: PH Website Company (Hilbert Magculang)\n"
        "Status: Root cause identified. Permanent fix built, tested and deployed September 5, 2026."
    ).font.size = Pt(10)
    meta.paragraph_format.space_after = Pt(12)

    # ---- 1. Summary ----
    heading(doc, "1. Summary")
    para(doc,
         "On the afternoon of September 4, six online bingo payments were successfully charged to customers' cards, "
         "but the booking website did not record a reservation for them. Two of those customers, receiving no "
         "confirmation, paid a second (and in one case a third) time. Staff learned of the problem when customers "
         "called the next day.")
    para(doc,
         "The direct cause was outside our system: our payment processor, Authorize.Net, sent the \"payment "
         "captured\" notification for these six payments roughly one hour late, while the notifications for the "
         "other sixteen payments that afternoon arrived within one to two minutes as normal. However, our booking "
         "system was designed in a way that depended entirely on that one notification arriving on time. When it "
         "did not, a safety rule held the money aside for manual review instead of completing the booking, and the "
         "alert for that review was not visible to staff.")
    para(doc,
         "No money was lost and no seat was double-sold. Every charge is accounted for and listed in section 4. "
         "We have changed the system so that a late or missing notification from the processor can no longer "
         "leave a paid customer without a seat, and so that staff are told immediately, with a one-click fix, if "
         "anything ever does need attention.", bold=False)

    # ---- 2. Customer experience ----
    heading(doc, "2. What the customer experienced")
    numbered(doc, [
        "Chose a seat (which placed a 20-minute hold on it), entered their name and email, and paid on the secure card form.",
        "The card was charged successfully by Authorize.Net.",
        "The website's page kept showing \"Confirming your payment...\" because it was waiting for the processor's confirmation, which did not come.",
        "After a while the customer gave up, or tried again from the start. Two customers were charged more than once as a result.",
        "About an hour after the payment, the processor's confirmation finally arrived. By then the 20-minute seat hold had expired, so the system did not attach the payment to the seat and placed it in a \"needs review\" state.",
    ])

    # ---- 3. Root cause ----
    heading(doc, "3. Why it happened")
    para(doc, "Three things combined. Any one of them alone would not have caused a problem.", italic=True)
    numbered(doc, [
        ("Late notifications from the payment processor. ",
         "Authorize.Net normally tells us within seconds that a payment succeeded. For six payments on September 4 that "
         "notice was delayed by 60 to 61 minutes. Our records show the website was running continuously (no restart, no "
         "outage) and that the other payments that afternoon were confirmed normally in between the delayed ones, so the "
         "delay was on the processor's side. We are raising this with Authorize.Net (see section 6)."),
        ("The website relied on that one notification. ",
         "There was no second way for the site to find out a payment had gone through. It could not ask the processor "
         "directly, and the card form's own \"success\" message back to our page was not reliably reaching us. So when "
         "the notification was late, the site simply waited."),
        ("The seat hold expired before the confirmation arrived. ",
         "To stop two people buying the same seat, a chosen seat is held for 20 minutes during checkout. When the "
         "payment confirmation finally arrived after the hold had lapsed, a safety rule refused to attach the payment "
         "to a seat that was no longer held, and parked the payment for staff review. That rule exists to prevent "
         "double-selling, but here the seats were still free, so the customer should simply have been given their seat. "
         "The review alert was only shown as a background message the admin screen did not display, so nobody saw it."),
    ])

    # ---- 4. Affected customers ----
    heading(doc, "4. Who was affected and what happens to each")
    para(doc, "All six payments were identified from the booking system's payment audit trail and cross-checked against "
              "Authorize.Net's captured transactions. Times are Atlantic Daylight Time.", size=10, color=GREY)
    table(doc,
          ["Customer", "Session", "Paid", "What happened", "Action"],
          [
              ["Dawn Fraser", "Sep 5, 6:00 pm", "$164.00",
               "Charged 3:28 pm; confirmation arrived 4:29 pm. No reservation recorded.",
               "Confirm her seat (Table 1, Chair 1) from the admin panel; call her to reassure."],
              ["Dan Johnson", "Nov 7, Big Bank Bingo", "$175.00",
               "Charged 1:25 pm; confirmation 2:26 pm. Staff manually held Table 42 Chair 5.",
               "Attach the $175 payment to that seat (remove the temporary promo ticket, then Confirm seat)."],
              ["Ashley Tucker", "Sep 5, 6:00 pm", "$84.00",
               "Charged 12:49 pm; confirmation 1:50 pm. Staff manually held Table 45 Chair 5.",
               "Confirm her paid seat (Table 49 Chair 4) and remove the temporary ticket, or keep 45/5 and refund $84. Tell her which."],
              ["Wendi Colton", "Sep 6, 6:30 pm", "$124.00 x 3",
               "Paid at 11:46 am and 12:17 pm with no confirmation; third payment at 2:34 pm confirmed normally. One seat (Table 54 Chair 4).",
               "Refund the two extra $124 charges."],
              ["Annie Arsenault", "Sep 4, 6:30 pm", "$300.00 x 2",
               "Paid 3:31 pm with no confirmation; paid again 3:54 pm, confirmed. First confirmation arrived 4:31 pm.",
               "Refund the duplicate $300 charge."],
              ["Bertha Dedam", "Aug 23, Bigger Bank Bingo", "$300.00",
               "Earlier, similar case (Aug 20): confirmation 3 hours late; seats had been sold on. She holds a separate paid booking.",
               "Confirm with her that the $300 was refunded; refund if not."],
          ],
          widths=[1.1, 1.15, 0.8, 2.35, 2.0])
    para(doc, "All refunds will be issued through the booking system's refund workflow, which reverses the charge with "
              "Authorize.Net and keeps the seat records consistent. Customers should expect refunds to appear on their "
              "statement within 5 to 10 business days of processing.", size=10)

    # ---- 5. What we fixed ----
    heading(doc, "5. What we changed so this cannot happen again")
    para(doc, "All changes were built, tested (automated regression tests covering every scenario from this incident, "
              "plus the full existing test suite) and deployed to the live booking site on September 5, 2026.")
    numbered(doc, [
        ("The site now asks the processor directly. ",
         "Every minute, and immediately while any customer is waiting on the confirmation page, the booking system "
         "checks Authorize.Net's own transaction list for our pending bookings and confirms any that were paid. "
         "Confirmation no longer depends on the processor's notification arriving on time, or at all."),
        ("A late payment keeps its seat. ",
         "If a confirmation arrives after the 20-minute hold has lapsed and the seat is still free, the paying customer "
         "gets the seat. The system only stops to ask for review when another customer is actively holding or has "
         "already bought that seat, which is the real double-selling risk."),
        ("The seat hold cannot expire while the customer is still on the page. ",
         "While a customer is on the card form or the confirming page, their seat hold is kept alive automatically "
         "(up to 90 minutes), so a slow processor cannot expire it under them."),
        ("Faster confirmation from the card form itself. ",
         "The moment the card form reports success, the browser now passes the transaction reference straight to our "
         "server, which verifies it with Authorize.Net and confirms the booking within seconds."),
        ("Staff are alerted, with a one-click fix. ",
         "If a payment ever does need attention, every supervisor account receives an email immediately, and the admin "
         "dashboard shows a red \"Payments Needing Attention\" panel listing the customer, seat, amount and what to do. "
         "Where the seat is still free, one click confirms the booking and emails the customer their ticket. "
         "Duplicate charges are listed for refund and can be marked as handled."),
        ("Clearer messages to customers. ",
         "The confirming page now tells customers their seat stays reserved and, if they were charged, not to pay again "
         "but to wait for the email or contact the bingo office with their booking reference."),
    ])

    # ---- 6. Follow-ups ----
    heading(doc, "6. Follow-up items")
    bullets(doc, [
        ("Authorize.Net: ", "we are asking the processor to explain the one-hour notification delay on September 4 "
                            "and to confirm the payment times for the affected transactions. Their explanation does not change our "
                            "fix; the site no longer depends on their notifications being on time."),
        ("Hosting configuration check: ", "we are confirming that the card form's return address on our hosting "
                                          "provider points to booking.wolastoqcasino.ca. If it does not, correcting it restores the fastest "
                                          "confirmation path for every customer. The new safeguards cover this either way."),
        ("Customer follow-up: ", "the six customers in section 4 should each be contacted by the bingo office once "
                                 "their seat is confirmed or refund is processed. Suggested wording is in section 7."),
        ("Monitoring: ", "the booking system now records every late confirmation and every recovered seat in its "
                         "audit log, and we will review those records weekly for the next month."),
    ])

    # ---- 7. Suggested customer message ----
    heading(doc, "7. Suggested message to affected customers")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(
        "\"Thank you for booking with Wolastoq Bingo, and we are sorry for the confusion with your online payment on "
        "September 4. Your payment was received, but a delay at our payment processor meant our website did not "
        "confirm your seat at the time. This has now been corrected: your seat for [session] is confirmed / your extra "
        "charge of [$amount] has been refunded and will appear on your statement within 5 to 10 business days. We have "
        "also updated the booking site so that this cannot happen again. If you have any questions, please call the "
        "bingo office and quote booking reference [BNG-...].\"")
    run.italic = True
    run.font.size = Pt(10.5)

    # ---- Appendix ----
    doc.add_page_break()
    heading(doc, "Appendix A. Technical timeline (UTC)")
    para(doc, "From the booking system's payment_events audit table. Each row is one booking; \"Delay\" is the time between "
              "the customer being sent to the card form and the processor's payment-captured notification reaching our server.",
         size=10, color=GREY)
    table(doc,
          ["Booking", "Customer", "Checkout started", "Notification received", "Delay", "System outcome"],
          [
              ["BNG-97A54FC584", "W. Colton", "14:46:27", "15:46:41", "60.2 min", "payment_review"],
              ["BNG-9A9599A389", "W. Colton", "15:17:31", "16:18:35", "61.1 min", "payment_review"],
              ["BNG-78F5673391", "A. Tucker", "15:49:46", "16:50:10", "60.4 min", "payment_review"],
              ["BNG-6560EE2314", "D. Johnson", "16:25:26", "17:26:05", "60.7 min", "payment_review"],
              ["BNG-5D570FAB1D", "D. Fraser", "17:28:45", "18:29:09", "60.4 min", "payment_review"],
              ["BNG-40BE8C8467", "A. Arsenault (1st charge)", "18:30:54", "19:31:42", "~60 min", "duplicate charge flagged"],
          ],
          widths=[1.3, 1.5, 1.1, 1.3, 0.75, 1.45])
    para(doc, "For comparison, the other 16 payments on September 4 were confirmed 0.2 to 2.6 minutes after checkout "
              "started. The web server had been running continuously for 56 hours with no restarts. Authorize.Net's "
              "published retry schedule for failed deliveries (3-minute, then 8-hour intervals) cannot produce a 60-minute "
              "gap, which places the delay in the processor's notification pipeline rather than in delivery to our server.",
         size=10)

    heading(doc, "Appendix B. Safeguards now in place", level=1)
    table(doc,
          ["Safeguard", "How it works", "Verified by"],
          [
              ["Gateway reconciliation", "Server polls Authorize.Net's unsettled (and, periodically, settled) transaction lists every 60 s and on demand; matches by booking reference; verifies each transaction server-to-server before confirming.", "payment-reconciliation-check"],
              ["Late-payment seat recovery", "Payment confirmed if all seats are free or held by this customer; review only if another customer actively holds or has bought the seat.", "late-payment-seat-reclaim-check"],
              ["Hold heartbeat", "Customer's status poll extends the seat hold while the checkout page is open; capped at 90 minutes.", "late-payment-seat-reclaim-check"],
              ["Direct transaction hand-off", "Card form success posts the transaction id to the server, which verifies it with the processor.", "late-payment-seat-reclaim-check"],
              ["Staff alerting and resolution", "Email to all supervisors plus dashboard panel with Confirm seat / Mark refunded actions.", "late-payment-seat-reclaim-check"],
              ["Existing protections retained", "Duplicate-charge detection, amount verification, seat double-sell prevention, refund approval workflow.", "Full regression suite (18 checks)"],
          ],
          widths=[1.6, 3.9, 1.9])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
